from __future__ import annotations

import html
import io
import re
from pathlib import Path
from typing import Any

from .db import EXPORT_ROOT, connect
from .paired_pdf_import import render_source_preview


def export_paper(paper: dict[str, Any], file_format: str, variant: str) -> Path:
    EXPORT_ROOT.mkdir(parents=True, exist_ok=True)
    safe_id = re.sub(r"[^a-zA-Z0-9_-]+", "-", paper["id"])
    target = EXPORT_ROOT / f"{safe_id}-{variant}.{file_format}"
    if file_format == "docx":
        export_docx(paper, variant, target)
    elif file_format == "pdf":
        export_pdf(paper, variant, target)
    else:
        raise ValueError("仅支持 pdf 和 docx 导出。")
    return target


def export_docx(paper: dict[str, Any], variant: str, target: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as error:
        raise RuntimeError("缺少 python-docx，请重新安装后端依赖。") from error

    document = Document()
    title = document.add_heading(paper["title"], level=0)
    title.alignment = 1
    meta = document.add_paragraph(f"模板：{paper.get('template', {}).get('name', '')}    版本：{variant}")
    meta.runs[0].font.size = Pt(10)

    questions = paper.get("questions", [])
    if not questions:
        document.add_paragraph("当前试卷没有选中题目，已生成空卷文件。")

    current_section = None
    question_number = 0
    for question in questions:
        section_title = question.get("section_title") or question.get("section_id") or "未命名分区"
        if section_title != current_section:
            current_section = section_title
            document.add_heading(section_title, level=1)
        question_number += 1
        question_image = source_preview(question, "question")
        document.add_paragraph(f"{question_number}.")
        if question_image:
            question_width, _ = fit_image_size(question_image, 6.3, 8.5)
            document.add_picture(io.BytesIO(question_image), width=Inches(question_width))
        else:
            document.add_paragraph("原题图片不可用，未导出识别文字。")
        if variant == "answer":
            analysis_image = source_preview(question, "analysis")
            if analysis_image:
                analysis_width, _ = fit_image_size(analysis_image, 6.3, 8.5)
                document.add_picture(io.BytesIO(analysis_image), width=Inches(analysis_width))
            else:
                answer = plain_formula_text(question.get("answer_markdown", ""))
                if answer:
                    document.add_paragraph(f"答案：{answer}")
                analysis = plain_formula_text(question.get("analysis_markdown", ""))
                if analysis:
                    document.add_paragraph(f"解析：{analysis}")
                points = question.get("scoring_points", [])
                if points:
                    document.add_paragraph(
                        "评分点：" + "；".join(f"{point.get('label', '')}（{point.get('score', 0)}分）" for point in points)
                    )
            if not analysis_image and not question.get("answer_markdown") and not question.get("analysis_markdown"):
                document.add_paragraph("暂无答案或解析。")
    document.save(target)


def export_pdf(paper: dict[str, Any], variant: str, target: Path) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as error:
        raise RuntimeError("缺少 reportlab，请重新安装后端依赖。") from error

    font_path = next(
        (
            path
            for path in (
                Path("C:/Windows/Fonts/msyh.ttc"),
                Path("C:/Windows/Fonts/simhei.ttf"),
                Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
            )
            if path.exists()
        ),
        None,
    )
    font_name = "Helvetica"
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("PaperHelperChinese", str(font_path)))
            font_name = "PaperHelperChinese"
        except Exception:
            font_name = "Helvetica"

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "PaperHelperBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=17,
        spaceAfter=5,
    )
    heading = ParagraphStyle(
        "PaperHelperHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        leading=20,
        spaceBefore=8,
        spaceAfter=8,
    )
    document = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=paper["title"],
    )
    flowables: list[Any] = [Paragraph(html.escape(paper["title"]), heading)]
    questions = paper.get("questions", [])
    if not questions:
        flowables.append(Paragraph("当前试卷没有选中题目，已生成空卷文件。", body))

    current_section = None
    question_number = 0
    for question in questions:
        section_title = question.get("section_title") or question.get("section_id") or "未命名分区"
        if section_title != current_section:
            current_section = section_title
            flowables.append(Paragraph(html.escape(section_title), heading))
        question_number += 1
        question_image = source_preview(question, "question")
        flowables.append(Paragraph(f"{question_number}.", body))
        if question_image:
            flowables.append(image_flowable(question_image))
        else:
            flowables.append(Paragraph("原题图片不可用，未导出识别文字。", body))
        if variant == "answer":
            analysis_image = source_preview(question, "analysis")
            if analysis_image:
                flowables.append(image_flowable(analysis_image))
            else:
                answer = question.get("answer_markdown", "")
                if answer:
                    flowables.append(Paragraph(f"<b>答案：</b>{safe_markup(answer)}", body))
                analysis = question.get("analysis_markdown", "")
                if analysis:
                    flowables.append(Paragraph(f"<b>解析：</b>{safe_markup(analysis)}", body))
                points = question.get("scoring_points", [])
                if points:
                    point_text = "；".join(
                        f"{point.get('label', '')}（{point.get('score', 0)}分）"
                        for point in points
                    )
                    flowables.append(Paragraph(f"<b>评分点：</b>{safe_markup(point_text)}", body))
            if not analysis_image and not question.get("answer_markdown") and not question.get("analysis_markdown"):
                flowables.append(Paragraph("暂无答案或解析。", body))
        flowables.append(Spacer(1, 3))
    document.build(flowables)


def export_practice_pdf(
    session: dict[str, Any],
    wrong_items: list[dict[str, Any]],
    target: Path,
) -> None:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as error:
        raise RuntimeError("缺少 reportlab，请重新安装后端依赖。") from error

    font_path = next(
        (
            path
            for path in (
                Path("C:/Windows/Fonts/msyh.ttc"),
                Path("C:/Windows/Fonts/simhei.ttf"),
                Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
            )
            if path.exists()
        ),
        None,
    )
    font_name = "Helvetica"
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("PracticeChinese", str(font_path)))
            font_name = "PracticeChinese"
        except Exception:
            font_name = "Helvetica"

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "PracticeBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=17,
        spaceAfter=5,
    )
    heading = ParagraphStyle(
        "PracticeHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=15,
        leading=21,
        spaceBefore=8,
        spaceAfter=8,
    )
    document = SimpleDocTemplate(
        str(target),
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        title=f"{session.get('subject', '')} 小题狂练错题",
    )
    range_text = (
        "范围：错题本全量抽取。"
        if session.get("practice_mode") == "wrong_book"
        else f"范围：{session.get('major_tag') or '全部大类'} / {session.get('sub_tag') or '全部小类'}。"
    )
    flowables: list[Any] = [
        Paragraph(html.escape(f"{session.get('subject', '')} 小题狂练错题"), heading),
        Paragraph(
            safe_markup(f"{range_text} 错题数：{len(wrong_items)}"),
            body,
        ),
        Spacer(1, 5),
    ]
    for index, item in enumerate(wrong_items, start=1):
        question = item["question"]
        attempt = item.get("attempt", {})
        flowables.append(Paragraph(f"{index}.", body))
        question_image = source_preview(question, "question")
        if question_image:
            flowables.append(image_flowable(question_image))
        else:
            flowables.append(Paragraph("原题图片不可用，未导出识别文字。", body))
        flowables.append(
            Paragraph(
                f"<b>你的答案：</b>{html.escape(', '.join(attempt.get('selected_options', [])))}"
                f"　<b>正确答案：</b>{html.escape(', '.join(attempt.get('correct_options', [])))}",
                body,
            )
        )
        analysis_image = source_preview(question, "analysis")
        if analysis_image:
            flowables.append(image_flowable(analysis_image))
        else:
            analysis = question.get("analysis_markdown", "")
            if analysis:
                flowables.append(Paragraph(f"<b>解析：</b>{safe_markup(analysis)}", body))
        flowables.append(Spacer(1, 6))
    document.build(flowables)


def plain_formula_text(value: str) -> str:
    return re.sub(r"(?<!\\)\$([^$]+)(?<!\\)\$", r"[\1]", value or "")


def safe_markup(value: str) -> str:
    text = html.escape(value or "")
    text = re.sub(r"\\?\\$([^$]+)\\?\\$", r"<font color='#2858ff'>[\1]</font>", text)
    text = text.replace("\n", "<br/>")
    return text


def source_preview(question: dict[str, Any], kind: str) -> bytes | None:
    source_id = (
        question.get("analysis_source_document_id")
        if kind == "analysis"
        else question.get("source_document_id")
    )
    regions = question.get("analysis_regions", []) if kind == "analysis" else question.get("source_regions", [])
    if kind == "question" and not regions and question.get("source_page"):
        regions = [{"page": question["source_page"], "bbox": [0.0, 0.0, 10000.0, 10000.0]}]
    if not source_id or not regions:
        return None
    with connect() as connection:
        source = connection.execute(
            "SELECT file_path FROM source_documents WHERE id = ?",
            (source_id,),
        ).fetchone()
    if source is None or not source["file_path"]:
        return None
    path = Path(source["file_path"])
    try:
        if not path.is_file():
            return None
    except OSError:
        return None
    try:
        return render_source_preview(path, regions, scale=1.35)
    except (OSError, RuntimeError, ValueError):
        return None


def image_flowable(content: bytes) -> Any:
    from reportlab.lib.units import mm
    from reportlab.platypus import Image

    display_width, display_height = fit_image_size(content, 174 * mm, 230 * mm)
    return Image(io.BytesIO(content), width=display_width, height=display_height)


def fit_image_size(content: bytes, max_width: float, max_height: float) -> tuple[float, float]:
    """Fit a rendered source image inside a page without changing its ratio."""
    try:
        from PIL import Image

        with Image.open(io.BytesIO(content)) as image:
            width, height = image.size
    except Exception as error:
        raise RuntimeError("无法读取原版题面图片尺寸。") from error
    scale = min(max_width / max(width, 1), max_height / max(height, 1), 1)
    return max(1.0, width * scale), max(1.0, height * scale)
