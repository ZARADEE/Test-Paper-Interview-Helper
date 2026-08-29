from __future__ import annotations

import hashlib
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from .math_one import classify_math_one_text


PDF_LAYOUT_NOISE = re.compile(
    r"^\s*第[一二三四五六七八九十百\d]+(?:章|部分)"
)


def _is_layout_noise(text: str) -> bool:
    compact = re.sub(r"\s+", "", text or "")
    if PDF_LAYOUT_NOISE.match(compact):
        return True
    return len(compact) <= 48 and any(
        label in compact
        for label in ("单项选择题", "多项选择题", "材料分析题", "参考答案")
    )


def _without_layout_noise(lines: list[dict[str, Any]]) -> list[dict[str, Any]]:
    cleaned: list[dict[str, Any]] = []
    noise_bands = [
        (line["bbox"][1] - 4.0, line["bbox"][3] + 28.0)
        for line in lines
        if _is_layout_noise(line["text"])
    ]
    for line in lines:
        y0, y1 = line["bbox"][1], line["bbox"][3]
        if any(y0 <= band_end and y1 >= band_start for band_start, band_end in noise_bands):
            continue
        cleaned.append(line)
    return cleaned


@dataclass
class DocumentPage:
    page_number: int
    text: str
    width: float | None = None
    height: float | None = None
    lines: list[dict[str, Any]] = field(default_factory=list)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def extract_document(path: Path) -> list[DocumentPage]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".doc":
        return extract_doc(path)
    raise ValueError("暂不支持该文件格式，请选择 PDF、DOCX 或 DOC。")


def extract_pdf(path: Path) -> list[DocumentPage]:
    try:
        import fitz
    except ImportError as error:
        raise RuntimeError("缺少 PyMuPDF，请重新安装后端依赖。") from error

    document = fitz.open(path)
    pages: list[DocumentPage] = []
    for index, page in enumerate(document):
        lines: list[dict[str, Any]] = []
        layout = page.get_text("dict", sort=True)
        for block in layout.get("blocks", []):
            for line in block.get("lines", []):
                text = "".join(span.get("text", "") for span in line.get("spans", []))
                if text.strip():
                    lines.append(
                        {
                            "text": text,
                            "bbox": [float(value) for value in line["bbox"]],
                        }
                    )
        pages.append(
            DocumentPage(
                index + 1,
                page.get_text("text"),
                float(page.rect.width),
                float(page.rect.height),
                lines,
            )
        )
    return pages


def extract_docx(path: Path) -> list[DocumentPage]:
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("缺少 python-docx，请重新安装后端依赖。") from error

    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = []
    for table in document.tables:
        for row in table.rows:
            table_text.append(" | ".join(cell.text.strip() for cell in row.cells))
    text = "\n".join([*paragraphs, *table_text])
    return [DocumentPage(1, text)]


def extract_doc(path: Path) -> list[DocumentPage]:
    libreoffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not libreoffice:
        raise RuntimeError("解析 DOC 需要 LibreOffice，请安装后重试。")

    with tempfile.TemporaryDirectory(prefix="paper-helper-doc-") as temp_dir:
        output_dir = Path(temp_dir)
        command = [
            libreoffice,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(output_dir),
            str(path),
        ]
        result = subprocess.run(command, capture_output=True, text=True, timeout=120)
        if result.returncode != 0:
            raise RuntimeError(result.stderr.strip() or "DOC 转换失败。")
        converted = output_dir / f"{path.stem}.docx"
        if not converted.exists():
            raise RuntimeError("LibreOffice 未生成转换后的 DOCX 文件。")
        return extract_docx(converted)


def split_question_candidates(pages: list[DocumentPage]) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    pattern = re.compile(r"(?m)^\s*(\d{1,3})\s*[.\u3001\uff0c,\u30fb)]\s*")
    for page in pages:
        matches = list(pattern.finditer(page.text))
        if not matches:
            text = page.text.strip()
            if text:
                candidates.append(
                    build_candidate(
                        text,
                        page.page_number,
                        0.45,
                        page.width,
                        page.height,
                    )
                )
            continue
        marker_regions = question_marker_regions(page, pattern)
        for index, match in enumerate(matches):
            start = match.start()
            end = matches[index + 1].start() if index + 1 < len(matches) else len(page.text)
            raw = page.text[start:end].strip()
            if raw:
                candidates.append(
                    build_candidate(
                        raw,
                        page.page_number,
                        0.78,
                        page.width,
                        page.height,
                        marker_regions[index] if index < len(marker_regions) else None,
                    )
                )
    return candidates


def question_marker_regions(
    page: DocumentPage,
    pattern: re.Pattern[str],
) -> list[list[dict[str, Any]]]:
    if not page.lines or not page.width or not page.height:
        return []
    lines = sorted(
        page.lines,
        key=lambda item: (round(item["bbox"][1], 2), round(item["bbox"][0], 2)),
    )
    markers = [
        (index, line)
        for index, line in enumerate(lines)
        if pattern.match(line["text"])
    ]
    if not markers:
        return []
    regions: list[list[dict[str, Any]]] = []
    for marker_index, (line_index, marker) in enumerate(markers):
        next_marker_y = (
            markers[marker_index + 1][1]["bbox"][1] - 3.0
            if marker_index + 1 < len(markers)
            else page.height - 18.0
        )
        segment_lines = _without_layout_noise(
            [
                line
                for line in lines
                if line["bbox"][1] >= marker["bbox"][1] - 2.0
                and line["bbox"][1] < next_marker_y + 2.0
            ]
        )
        if not segment_lines:
            return []
        if marker_index + 1 == len(markers):
            segment_lines = [
                line
                for line in segment_lines
                if line["bbox"][1] < page.height - 24.0
            ] or segment_lines
        x0 = max(0.0, min(line["bbox"][0] for line in segment_lines) - 8.0)
        y0 = max(0.0, min(line["bbox"][1] for line in segment_lines) - 6.0)
        x1 = min(page.width, max(line["bbox"][2] for line in segment_lines) + 8.0)
        y1 = min(page.height, max(line["bbox"][3] for line in segment_lines) + 6.0)
        if y1 <= y0 + 2:
            return []
        regions.append(
            [
                {
                    "page": page.page_number,
                    "bbox": [x0, y0, x1, y1],
                }
            ]
        )
    return regions


def build_candidate(
    raw_text: str,
    page_number: int,
    confidence: float,
    page_width: float | None = None,
    page_height: float | None = None,
    source_regions: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    classified = classify_math_one_text(raw_text)
    math_signal = any(
        keyword in raw_text
        for keyword in ("极限", "导数", "积分", "矩阵", "行列式", "随机变量", "概率", "方程")
    )
    if not math_signal:
        classified.update(
            {
                "subject": "待分类",
                "tags": [],
                "chapter": "",
                "knowledge_points": [],
                "confidence": min(confidence, 0.45),
            }
        )
    else:
        classified["confidence"] = min(confidence, float(classified.get("confidence", confidence)))
    resolved_regions = source_regions or []
    if not resolved_regions and page_width and page_height:
        resolved_regions = [
            {
                "page": page_number,
                "bbox": [0.0, 0.0, page_width, page_height],
            }
        ]
    return {
        **classified,
        "stem_markdown": raw_text,
        "answer_markdown": "",
        "analysis_markdown": "",
        "scoring_points": [],
        "source_page": page_number,
        "source_regions": resolved_regions,
    }
